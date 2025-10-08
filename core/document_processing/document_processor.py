"""
Document Processor - Handles document processing and uploads to Pinecone.
Integrates with Pinecone's automatic embedding service.
"""

import os
import logging
import hashlib
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore
import pdfplumber
from docx import Document
from openpyxl import load_workbook
import pandas as pd
from bs4 import BeautifulSoup
from fastapi import UploadFile
from sqlalchemy.orm import Session

from core.document_processing.text_splitter import TextSplitter
from core.document_processing.file_processor import FileProcessor
from core.pinecone.pinecone_service import PineconeService
from core.database.models import Document as DBDocument, DocumentType, Department
from core.llm.config import get_settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Process documents and upload to Pinecone with integrated embeddings.
    """
    
    def __init__(
        self,
        pinecone_service: PineconeService,
        db: Session,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        """
        Initialize Document Processor.
        
        Args:
            pinecone_service: Pinecone service instance
            db: Database session
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.pinecone_service = pinecone_service
        self.db = db
        self.text_splitter = TextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.file_processor = FileProcessor()
        self.settings = get_settings()
        
        logger.info(f"DocumentProcessor initialized with chunk_size={chunk_size}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using multiple methods.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        text = ""
        
        # Try PyMuPDF first (faster)
        if fitz is not None:
            try:
                with fitz.open(file_path) as doc:  # type: ignore
                    for page in doc:
                        text += page.get_text()
                
                if text.strip():
                    logger.info(f"Extracted {len(text)} characters from PDF using PyMuPDF")
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}, trying pdfplumber...")
        
        # Fallback to pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF using pdfplumber")
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise
        
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info(f"Extracted {len(text)} characters from TXT")
            return text
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Extracted {len(text)} characters from TXT using {encoding}")
                    return text
                except:
                    continue
            raise ValueError("Could not decode text file with any known encoding")
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file."""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            text = ""
            
            for sheet_name, sheet_df in df.items():
                text += f"\n\n=== Sheet: {sheet_name} ===\n"
                # Convert DataFrame to text representation
                text += sheet_df.to_string(index=False)
            
            logger.info(f"Extracted {len(text)} characters from Excel")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from Excel: {e}")
            raise
    
    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            logger.info(f"Extracted {len(text)} characters from HTML")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from HTML: {e}")
            raise
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from file based on extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        extractors = {
            '.pdf': self.extract_text_from_pdf,
            '.docx': self.extract_text_from_docx,
            '.txt': self.extract_text_from_txt,
            '.xlsx': self.extract_text_from_excel,
            '.xls': self.extract_text_from_excel,
            '.csv': self.extract_text_from_excel,
            '.html': self.extract_text_from_html,
            '.htm': self.extract_text_from_html,
        }
        
        extractor = extractors.get(ext)
        if not extractor:
            raise ValueError(f"Unsupported file type: {ext}")
        
        return extractor(file_path)
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def create_document_metadata(
        self,
        file_name: str,
        file_path: str,
        document_type: str,
        chunk_id: int,
        total_chunks: int,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Create metadata for a document chunk.
        
        Args:
            file_name: Original filename
            file_path: Path to file
            document_type: Type of document
            chunk_id: Current chunk index
            total_chunks: Total number of chunks
            additional_metadata: Additional metadata to include
            
        Returns:
            Metadata dictionary
        """
        metadata = {
            "original_filename": file_name,
            "file_name": file_name,
            "document_type": document_type,
            "chunk_id": chunk_id,
            "total_chunks": total_chunks,
            "upload_date": datetime.now().isoformat(),
            "file_size": os.path.getsize(file_path),
            "file_hash": self.calculate_file_hash(file_path),
        }
        
        if additional_metadata:
            metadata.update(additional_metadata)
        
        return metadata
    
    async def process_and_upload_file(
        self,
        file: UploadFile,
        namespace: str = "default",
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process a file and upload to Pinecone.
        
        Args:
            file: Uploaded file
            namespace: Pinecone namespace
            additional_metadata: Additional metadata
            
        Returns:
            Dictionary with upload results
        """
        # Save uploaded file temporarily
        temp_dir = self.settings.UPLOAD_DIR
        os.makedirs(temp_dir, exist_ok=True)
        
        # Ensure filename is not None
        filename = file.filename or "unnamed_file"
        file_path = os.path.join(temp_dir, filename)
        
        try:
            # Save file
            with open(file_path, 'wb') as f:
                content = await file.read()
                f.write(content)
            
            logger.info(f"Processing file: {filename}")
            
            # Extract text
            text = self.extract_text_from_file(file_path)
            
            if not text or len(text.strip()) < 10:
                raise ValueError("Extracted text is too short or empty")
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Split into {len(chunks)} chunks")
            
            # Prepare documents for Pinecone
            documents = []
            base_doc_id = str(uuid.uuid4())
            
            for i, chunk_text in enumerate(chunks):
                doc_id = f"{base_doc_id}_chunk_{i}"
                
                # Create flat metadata structure (Pinecone v7 requirement)
                # All fields must be at top level, not nested
                doc_metadata = {
                    "source": filename,
                    "document_type": self.file_processor.get_file_type(filename),
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "upload_date": datetime.now().strftime("%Y-%m-%d"),
                    "file_hash": self.calculate_file_hash(file_path)[:16]  # Shortened for metadata
                }
                
                # Add additional metadata if provided (keep flat)
                if additional_metadata:
                    for key, value in additional_metadata.items():
                        # Only add simple types (str, int, float, bool)
                        if isinstance(value, (str, int, float, bool)):
                            doc_metadata[key] = value
                
                # Format for Pinecone v7 (auto-embedding via integrated inference)
                documents.append({
                    "id": doc_id,
                    "chunk_text": chunk_text,  # This field gets auto-embedded by Pinecone
                    **doc_metadata  # Flat metadata at top level
                })
            
            # Upload to Pinecone
            logger.info(f"Uploading {len(documents)} chunks to Pinecone...")
            upload_result = self.pinecone_service.upsert_documents(
                documents=documents,
                namespace=namespace
            )
            
            # Save to PostgreSQL/SQLite
            db_document = DBDocument(
                document_id=base_doc_id,
                file_name=filename,
                display_name=filename,
                file_type=self.file_processor.get_file_type(filename),
                file_size=os.path.getsize(file_path),
                file_hash=self.calculate_file_hash(file_path),
                total_chunks=len(chunks)
                # created_at is auto-set by model
                # namespace is stored in Pinecone, not in DB
            )
            
            if additional_metadata:
                if 'document_type' in additional_metadata:
                    db_document.document_type = additional_metadata['document_type']
                if 'department' in additional_metadata:
                    db_document.department = additional_metadata['department']
            
            self.db.add(db_document)
            self.db.commit()
            
            logger.info(f"Successfully processed and uploaded: {filename}")
            
            return {
                "status": "success",
                "filename": filename,
                "document_id": base_doc_id,
                "chunks_count": len(chunks),
                "namespace": namespace,
                "pinecone_upload": upload_result
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise
        
        finally:
            # Clean up temporary file
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug(f"Cleaned up temp file: {file_path}")
    
    async def process_and_upload_file_from_path(
        self,
        file_path: str,
        original_filename: str,
        namespace: str = "default",
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process and upload a file from disk path (for background tasks).
        Similar to process_and_upload_file but reads from file path instead of UploadFile.
        
        Args:
            file_path: Path to the saved file
            original_filename: Original filename
            namespace: Pinecone namespace
            additional_metadata: Additional metadata to include
            
        Returns:
            Status dict with upload results
        """
        try:
            # Extract text from file
            text = self.extract_text_from_file(file_path)
            logger.info(f"Extracted {len(text)} characters from {original_filename}")
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Split into {len(chunks)} chunks")
            
            # Generate base document ID
            base_doc_id = str(uuid.uuid4())
            
            # Prepare documents with flat metadata for Pinecone v7
            documents = []
            for i, chunk_text in enumerate(chunks):
                doc_id = f"{base_doc_id}_chunk_{i}"
                
                # Create flat metadata structure (Pinecone v7 requirement)
                doc_metadata = {
                    "source": original_filename,
                    "document_type": self.file_processor.get_file_type(original_filename),
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "upload_date": datetime.now().strftime("%Y-%m-%d"),
                    "file_hash": self.calculate_file_hash(file_path)[:16]
                }
                
                # Add additional metadata (keep flat, only simple types)
                if additional_metadata:
                    for key, value in additional_metadata.items():
                        if isinstance(value, (str, int, float, bool)):
                            doc_metadata[key] = value
                
                documents.append({
                    "id": doc_id,
                    "chunk_text": chunk_text,  # REQUIRED field for Pinecone
                    **doc_metadata  # Flat at top level
                })
            
            # Upload to Pinecone (auto-embedding handled by Pinecone)
            upload_result = self.pinecone_service.upsert_documents(
                documents=documents,
                namespace=namespace
            )
            
            # Save to Database (SQLite)
            db_document = DBDocument(
                document_id=base_doc_id,
                file_name=original_filename,
                display_name=original_filename,
                file_type=self.file_processor.get_file_type(original_filename),
                file_size=os.path.getsize(file_path),
                file_hash=self.calculate_file_hash(file_path),
                total_chunks=len(chunks)
                # created_at is auto-set by model
                # namespace is stored in Pinecone metadata, not in DB
            )
            
            if additional_metadata:
                if 'document_type' in additional_metadata:
                    db_document.document_type = additional_metadata['document_type']
                if 'department' in additional_metadata:
                    db_document.department = additional_metadata['department']
            
            self.db.add(db_document)
            self.db.commit()
            
            logger.info(f"Successfully processed and uploaded: {original_filename}")
            
            return {
                "status": "success",
                "filename": original_filename,
                "document_id": base_doc_id,
                "chunks_count": len(chunks),
                "namespace": namespace,
                "pinecone_upload": upload_result
            }
            
        except Exception as e:
            logger.error(f"Error processing file {original_filename}: {str(e)}")
            return {
                "status": "error",
                "error": str(e),
                "filename": original_filename
            }
    
    async def delete_document(
        self,
        document_id: str,
        namespace: str = "default"
    ) -> bool:
        """
        Delete a document from both Pinecone and PostgreSQL.
        
        Args:
            document_id: Base document ID
            namespace: Pinecone namespace
            
        Returns:
            True if successful
        """
        try:
            # Get document from DB to find chunk count
            db_doc = self.db.query(DBDocument).filter(
                DBDocument.document_id == document_id
            ).first()
            
            if not db_doc:
                logger.warning(f"Document {document_id} not found in database")
                return False
            
            # Generate all chunk IDs
            chunk_ids = [
                f"{document_id}_chunk_{i}"
                for i in range(db_doc.chunk_count)
            ]
            
            # Delete from Pinecone
            self.pinecone_service.delete_vectors(chunk_ids, namespace)
            
            # Delete from PostgreSQL
            self.db.delete(db_doc)
            self.db.commit()
            
            logger.info(f"Successfully deleted document: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            self.db.rollback()
            raise
