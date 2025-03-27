import os
import logging
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional, Union
import fitz
import pdfplumber
import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document
from openpyxl import load_workbook
import pandas as pd
from bs4 import BeautifulSoup
import tabula
import json
from enum import Enum
import numpy as np

# Configure logging
logger = logging.getLogger(__name__)

class ChunkingStrategy(Enum):
    """Defines different strategies for chunking tabular data."""
    NO_CHUNK = "no_chunk"  # Keep all records together
    FIXED_SIZE = "fixed_size"  # Split into fixed-size chunks
    SMART_CHUNK = "smart_chunk"  # Use smart chunking based on data characteristics

class TabularDataProcessor:
    """Handles processing and structuring of tabular data from various sources."""
    
    def __init__(
        self,
        chunk_size: int = 100,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.SMART_CHUNK,
        min_chunk_size: int = 50,
        max_chunk_size: int = 500
    ):
        """
        Initialize TabularDataProcessor.
        
        Args:
            chunk_size: Target number of records per chunk for table data
            chunking_strategy: Strategy to use for chunking table data
            min_chunk_size: Minimum records per chunk for SMART_CHUNK strategy
            max_chunk_size: Maximum records per chunk for SMART_CHUNK strategy
        """
        self.chunk_size = chunk_size
        self.chunking_strategy = chunking_strategy
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
    
    def _record_to_flat_text(self, record: Dict[str, Any]) -> str:
        """
        Convert a record dictionary to flat text format including field names.
        
        Args:
            record: Dictionary containing the record data
            
        Returns:
            A flat text representation of the record with field names
        """
        try:
            # Convert special types to string with field names
            text_parts = []
            for key, value in record.items():
                if pd.isna(value):
                    text_parts.append(f"{key}: ")
                elif isinstance(value, pd.Timestamp):
                    text_parts.append(f"{key}: {value.strftime('%Y-%m-%d %H:%M:%S')}")
                elif isinstance(value, (pd.Series, pd.DataFrame)):
                    text_parts.append(f"{key}: {str(value.to_dict())}")
                elif isinstance(value, np.integer):
                    text_parts.append(f"{key}: {str(int(value))}")
                elif isinstance(value, np.floating):
                    text_parts.append(f"{key}: {str(float(value))}")
                else:
                    text_parts.append(f"{key}: {str(value)}")
            
            # Join all field-value pairs with commas
            return ", ".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error converting record to flat text: {str(e)}")
            return str(record)

    def process_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """Process Excel files and extract structured table data."""
        try:
            tables = []
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            for sheet_name, df in excel_data.items():
                if not df.empty:
                    # Clean column names
                    df.columns = [str(col).strip() if str(col).strip() else f"Column_{i+1}" 
                                for i, col in enumerate(df.columns)]
                    
                    # Convert each record to flat text
                    records = df.replace({pd.NA: None}).to_dict('records')
                    for record in records:
                        # Convert record to flat text
                        description = self._record_to_flat_text(record)
                        
                        table_data = {
                            "description": description,
                            "original_record": record
                        }
                        tables.append(table_data)
            
            return tables
            
        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            return []

    def process_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """Process CSV files and convert records to flat text."""
        try:
            tables = []
            # Read CSV in chunks to handle large files
            for chunk in pd.read_csv(file_path, chunksize=self.chunk_size):
                if not chunk.empty:
                    # Clean column names
                    chunk.columns = [str(col).strip() if str(col).strip() else f"Column_{i+1}" 
                                   for i, col in enumerate(chunk.columns)]
                    
                    # Process each record
                    records = chunk.replace({pd.NA: None}).to_dict('records')
                    for record in records:
                        # Convert record to flat text
                        description = self._record_to_flat_text(record)
                        
                        table_data = {
                            "description": description,
                            "original_record": record
                        }
                        tables.append(table_data)
            
            return tables
            
        except Exception as e:
            logger.error(f"Error processing CSV file: {str(e)}")
            return []

    def process_pdf_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract tables from PDF files."""
        try:
            tables = []
            # Try lattice mode first
            pdf_tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True, 
                                       lattice=True, stream=False, silent=True)
            
            # If no tables found, try stream mode
            if not pdf_tables:
                pdf_tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True,
                                           lattice=False, stream=True, silent=True)
            
            for page_num, df in enumerate(pdf_tables):
                if isinstance(df, pd.DataFrame) and not df.empty:
                    # Clean column names
                    df.columns = [str(col).strip() if str(col).strip() else f"Column_{i+1}" 
                                for i, col in enumerate(df.columns)]
                    
                    # Convert DataFrame to records and chunk them
                    chunks = self._chunk_dataframe(df)
                    
                    for chunk_idx, chunk_records in enumerate(chunks):
                        table_data = {
                            "type": "pdf_table",
                            "page_number": page_num + 1,
                            "chunk_index": chunk_idx,
                            "total_chunks": len(chunks),
                            "headers": list(df.columns),
                            "records": chunk_records,
                            "row_count": len(chunk_records),
                            "column_count": len(df.columns)
                        }
                        tables.append(table_data)
            
            return tables
            
        except Exception as e:
            logger.error(f"Error extracting tables from PDF: {str(e)}")
            return []
    
    def process_docx_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract tables from Word documents."""
        try:
            tables = []
            doc = Document(file_path)
            
            for table_num, table in enumerate(doc.tables):
                data = []
                headers = []
                
                # Process headers
                if table.rows:
                    header_row = table.rows[0]
                    headers = [cell.text.strip() if cell.text.strip() else f"Column_{i+1}"
                             for i, cell in enumerate(header_row.cells)]
                
                # Process data rows
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data) and len(row_data) == len(headers):
                        data.append(row_data)
                
                if headers and data:
                    # Convert to DataFrame for consistent processing
                    df = pd.DataFrame(data, columns=headers)
                    chunks = self._chunk_dataframe(df)
                    
                    for chunk_idx, chunk_records in enumerate(chunks):
                        table_data = {
                            "type": "docx_table",
                            "table_number": table_num + 1,
                            "chunk_index": chunk_idx,
                            "total_chunks": len(chunks),
                            "headers": headers,
                            "records": chunk_records,
                            "row_count": len(chunk_records),
                            "column_count": len(headers)
                        }
                        tables.append(table_data)
            
            return tables
            
        except Exception as e:
            logger.error(f"Error extracting tables from Word document: {str(e)}")
            return []

class DocumentProcessor:
    """
    Handles document processing for various file types, including:
    - Text extraction
    - Document chunking with configurable overlap
    - Metadata extraction and management
    - Tabular data processing
    """
    
    SUPPORTED_EXTENSIONS = {
        'pdf': ['pdf'],
        'document': ['doc', 'docx', 'txt', 'rtf'],
        'presentation': ['ppt', 'pptx'],
        'spreadsheet': ['xls', 'xlsx', 'csv'],
        'web': ['html', 'htm'],
    }
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        upload_dir: str = "data/uploads",
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.upload_dir = upload_dir
        
        # Initialize tabular processor
        self.tabular_processor = TabularDataProcessor()
        
        # Create upload directory if it doesn't exist
        os.makedirs(self.upload_dir, exist_ok=True)
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )
    
    def process_file(self, file_path: str) -> Tuple[List[Union[str, Dict[str, Any]]], List[Dict[str, Any]]]:
        """
        Process a file and return its chunks and associated metadata.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Tuple containing:
                - List of text chunks (for text files) or descriptions (for tabular data)
                - List of metadata dictionaries for each chunk
        """
        try:
            # Extract file extension
            file_ext = os.path.splitext(file_path)[1].lower().replace('.', '')
            
            # Validate file type
            if not self._is_supported_file_type(file_ext):
                logger.error(f"Unsupported file type: {file_ext}")
                return [], []
            
            # Generate file hash for identification
            file_hash = self._generate_file_hash(file_path)
            
            # Extract basic metadata
            base_metadata = self._extract_metadata(file_path, file_ext)
            base_metadata["file_hash"] = file_hash
            
            content = []
            metadata_list = []
            
            # Process based on file type
            if file_ext in self.SUPPORTED_EXTENSIONS['spreadsheet']:
                # For spreadsheet files, process as tabular data
                if file_ext in ['xlsx', 'xls']:
                    tables = self.tabular_processor.process_excel(file_path)
                elif file_ext == 'csv':
                    tables = self.tabular_processor.process_csv(file_path)
                
                # Each table record becomes a separate content piece
                for table in tables:
                    # Only store the description as content
                    content.append(table["description"])
                    
                    # Simplified metadata
                    record_metadata = base_metadata.copy()
                    record_metadata.update({
                        "content_type": "table_record"
                    })
                    metadata_list.append(record_metadata)
            
            else:
                # For non-spreadsheet files, process text and embedded tables
                
                # Extract text content
                text = self._extract_text(file_path, file_ext)
                
                # Process any tables in the document
                tables = self._extract_tables(file_path, file_ext)
                
                # Add table records first
                for table in tables:
                    # Only store the description as content
                    content.append(table["description"])
                    
                    # Simplified metadata
                    record_metadata = base_metadata.copy()
                    record_metadata.update({
                        "content_type": "table_record"
                    })
                    metadata_list.append(record_metadata)
                
                # Then add text chunks
                if text:
                    chunks = self.text_splitter.split_text(text)
                    for i, chunk in enumerate(chunks):
                        content.append(chunk)
                        chunk_metadata = base_metadata.copy()
                        chunk_metadata.update({
                            "content_type": "text",
                            "chunk_id": i,
                            "total_chunks": len(chunks)
                        })
                        metadata_list.append(chunk_metadata)
            
            logger.info(f"Processed {file_path} into {len(content)} content pieces")
            return content, metadata_list
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return [], []
    
    def _extract_tables(self, file_path: str, file_ext: str) -> List[Dict[str, Any]]:
        """Extract tables from various file types."""
        try:
            if file_ext in ['xlsx', 'xls']:
                return self.tabular_processor.process_excel(file_path)
            elif file_ext == 'pdf':
                return self.tabular_processor.process_pdf_tables(file_path)
            elif file_ext in ['html', 'htm']:
                return self.tabular_processor.process_html_tables(file_path)
            elif file_ext in ['doc', 'docx']:
                return self.tabular_processor.process_docx_tables(file_path)
            else:
                return []
        except Exception as e:
            logger.error(f"Error extracting tables: {str(e)}")
            return []
    
    def _extract_text(self, file_path: str, file_ext: str) -> str:
        """Extract text from a file based on its extension."""
        try:
            if file_ext in self.SUPPORTED_EXTENSIONS['pdf']:
                return self._extract_pdf_text(file_path)
            elif file_ext in self.SUPPORTED_EXTENSIONS['document']:
                return self._extract_document_text(file_path, file_ext)
            elif file_ext in self.SUPPORTED_EXTENSIONS['spreadsheet']:
                return self._extract_spreadsheet_text(file_path, file_ext)
            elif file_ext in self.SUPPORTED_EXTENSIONS['web']:
                return self._extract_html_text(file_path)
            else:
                # For unsupported types, return empty string
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files using multiple libraries for better results."""
        try:
            # Try PyMuPDF first (better for most PDFs)
            text = ""
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text() + "\n\n"
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {str(e)}. Trying pdfplumber.")
                
            # If PyMuPDF fails or text is too short, try pdfplumber
            if len(text) < 100:
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
            
            # Clean up the text
            text = self._clean_text(text)
            return text
        
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {str(e)}")
            return ""
    
    def _extract_document_text(self, file_path: str, file_ext: str) -> str:
        """Extract text from document files (txt, doc, docx, rtf)."""
        try:
            if file_ext == 'txt':
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    # Try with different encoding if utf-8 fails
                    with open(file_path, 'r', encoding='latin-1') as file:
                        text = file.read()
                    return self._clean_text(text)
            elif file_ext in ['doc', 'docx']:
                doc = Document(file_path)
                text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
                return self._clean_text(text)
            else:
                logger.warning(f"Document type {file_ext} not fully supported yet")
                return ""
        except Exception as e:
            logger.error(f"Error extracting document text: {str(e)}")
            return ""
    
    def _extract_spreadsheet_text(self, file_path: str, file_ext: str) -> str:
        """Extract text from spreadsheet files (xlsx, xls, csv)."""
        try:
            if file_ext in ['xlsx', 'xls']:
                workbook = load_workbook(filename=file_path, read_only=True)
                text_parts = []
                
                for sheet in workbook.sheetnames:
                    worksheet = workbook[sheet]
                    sheet_text = []
                    sheet_text.append(f"Sheet: {sheet}")
                    
                    for row in worksheet.iter_rows(values_only=True):
                        # Filter out None values and convert all to strings
                        row_text = [str(cell) for cell in row if cell is not None]
                        if row_text:  # Only add non-empty rows
                            sheet_text.append("\t".join(row_text))
                    
                    text_parts.append("\n".join(sheet_text))
                
                workbook.close()
                return self._clean_text("\n\n".join(text_parts))
            elif file_ext == 'csv':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return self._clean_text(file.read())
            else:
                logger.warning(f"Spreadsheet type {file_ext} not fully supported yet")
                return ""
        except Exception as e:
            logger.error(f"Error extracting spreadsheet text: {str(e)}")
            return ""
    
    def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Simple HTML tag removal - can be improved with BeautifulSoup
            text = re.sub(r'<[^>]+>', ' ', html_content)
            text = re.sub(r'\s+', ' ', text)
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting HTML text: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove non-printable characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    def _is_supported_file_type(self, file_ext: str) -> bool:
        """Check if the file extension is supported."""
        for category, extensions in self.SUPPORTED_EXTENSIONS.items():
            if file_ext in extensions:
                return True
        return False
    
    def _generate_file_hash(self, file_path: str) -> str:
        """Generate a hash for the file to use as a unique identifier."""
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as file:
                # Read file in chunks to handle large files
                for chunk in iter(lambda: file.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"Error generating file hash: {str(e)}")
            # Return a timestamp-based fallback hash
            return f"hash_{int(datetime.now().timestamp())}"
    
    def _extract_metadata(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        """Extract metadata from the file."""
        try:
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            creation_time = os.path.getctime(file_path)
            modification_time = os.path.getmtime(file_path)
            
            metadata = {
                "file_name": file_name,
                "file_type": file_ext,
                "file_size": file_size,
                "creation_date": datetime.fromtimestamp(creation_time).isoformat(),
                "modification_date": datetime.fromtimestamp(modification_time).isoformat(),
                "source_path": file_path,
            }
            
            # Add category based on file extension
            for category, extensions in self.SUPPORTED_EXTENSIONS.items():
                if file_ext in extensions:
                    metadata["document_type"] = category
                    break
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {
                "file_name": os.path.basename(file_path),
                "file_type": file_ext,
                "error": str(e)
            } 